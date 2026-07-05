import os
import sys
from dotenv import load_dotenv
from google import genai
from google.genai import types
from docx import Document
import time

def evaluate_chatbot(input_file):
    with open(input_file,"r") as f:
        document = f.readlines()
    
    testcase = []
    expected = []
    forbidden = []
    results = []
    for line in document:
        if line.startswith("CHATBOT_NAME:"):
            chatbot_names = line.split("CHATBOT_NAME:", 1)
            chatbot_name = chatbot_names[1].strip()
            print(f"chatbotname: {chatbot_name}")
        if line.startswith("CHATBOT_CONTEXT:"):
            chatbot_context = line.split("CHATBOT_CONTEXT:", 1)
            chatbot_context = chatbot_context[1].strip()
            print(f"chatbotcontext: {chatbot_context}")
        if line.strip().startswith("TC") and "|" in line:
            cells = line.split("|")
            cells = [c.strip() for c in cells]
            testcase.append(cells[1])
            expected.append(cells[2])
            forbidden.append(cells[3])

    client = genai.Client(api_key=os.getenv("GEMINI_API_KEY"))

    for i,tc in enumerate(testcase):
        chatbot_prompt = f"""
            You are {chatbot_name}.
            {chatbot_context}
            User asks: {tc}
            Respond naturally as this chatbot.
        """
        chatbot_response = call_ai_with_retry(client,chatbot_prompt)
        time.sleep(3)  # Add a delay to avoid hitting rate limits


        evaluator_prompt = f"""
            You are a QA engineer evaluating a chatbot response.
            Test Case: {tc}
            Chatbot Response: {chatbot_response}
            Expected: {expected[i]}
            Should NOT contain: {forbidden[i]}
            Give PASS or FAIL with a reason.
            """
        
        qa_evaluation = call_ai_with_retry(client,evaluator_prompt)
        time.sleep(3)  # Add a delay to avoid hitting rate limits

        results.append({
            "tc_number": i + 1,
            "prompt": tc,
            "expected": expected[i],
            "forbidden": forbidden[i],
            "chatbot_response": chatbot_response,
            "evaluation": qa_evaluation
        })

    return results

def save_report(results, output_file):
    doc = Document()
    doc.add_heading("Chatbot QA Evaluation Report", level=0)

    #count the number of PASS and FAIL
    pass_count = sum(1 for result in results if "PASS" in result["evaluation"].upper())
    fail_count = len(results) - pass_count
    score = (pass_count / len(results)) * 100 if results else 0
    print(f"Evaluation Summary - PASS: {pass_count}, FAIL: {fail_count}")
    print(f"Overall Score: {score:.2f}%")

    doc.add_heading("Summary", level=1)
    doc.add_paragraph(f"Total Test Cases: {len(results)}")
    doc.add_paragraph(f"Passed: {pass_count}")
    doc.add_paragraph(f"Failed: {fail_count}")
    doc.add_paragraph(f"Quality Score: {score:.2f}%")

    doc.add_heading("Detailed Results", level=1)
    for result in results:
        doc.add_heading(f"Test Case {result['tc_number']}", level=2)
        doc.add_paragraph(f"Prompt: {result['prompt']}")
        doc.add_paragraph(f"Expected: {result['expected']}")
        doc.add_paragraph(f"Should NOT contain: {result['forbidden']}")
        doc.add_paragraph(f"Chatbot Response: {result['chatbot_response']}")
        doc.add_paragraph(f"Evaluation: {result['evaluation']}")
    
    doc.add_heading("Notes for QA Engineer", level=1)
    doc.add_paragraph(
        "This report was generated automatically using AI evaluation. "
        "All FAIL cases should be manually verified before raising defects. "
        "PASS cases should be spot-checked to ensure evaluation accuracy. "
        "AI is your assistant — you are still the engineer."
)

    doc.save(output_file)
    print(f"✅ Evaluation report saved to {output_file}")

MODELS = ["gemini-2.5-flash", "gemini-2.5-flash-lite"]

def call_ai_with_retry(client, prompt, retries=3):
    for model in MODELS:
        for attempt in range(retries):
            try:
                api_response = client.models.generate_content(
                    model=model,
                    contents=prompt,
                    config=types.GenerateContentConfig(temperature=0.1))
                print(f"✅ Using model: {model}")
                return api_response.text
            except Exception as e:
                if "503" in str(e) or "UNAVAILABLE" in str(e):
                    print(f"⚠️ {model} busy, retrying in 30 seconds... (attempt {attempt+1}/{retries})")
                    time.sleep(30)
                elif "429" in str(e) or "RESOURCE_EXHAUSTED" in str(e):
                    print(f"⚠️ {model} quota exceeded, trying next model...")
                    break
                else:
                    raise e
    print("❌ All models exhausted. Please try again tomorrow.")
    sys.exit(1)

if __name__ == "__main__":
    load_dotenv()
    if len(sys.argv) < 2: 
        print("Usage: python evaluate_chatbot.py inputs/testcases.txt")
        sys.exit(1)
    input_file = sys.argv[1]
    results = evaluate_chatbot(input_file)
    save_report(results, "outputs/chatbot_evaluation_report.docx")
